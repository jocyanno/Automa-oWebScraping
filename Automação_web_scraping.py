# JOCYANNO VITTOR
# @jocyannovittor@hotmail.com

import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import RGBColor
import schedule
import time
import win32com.client as win32

def webscrabing():

  lista_noticias = []

  response = requests.get('https://g1.globo.com/')

  content = response.content

  site = BeautifulSoup(content, 'html.parser')

  noticias = site.findAll('div', attrs={'class': 'feed-post-body'})

def fazerColocarNoticiasNoWord(noticias):
      
  documento = Document()

  arquivo = open('NoticiasDeHoje.doc','w')

  for noticia in noticias:

    titulo = noticia.find('a', attrs={'class': 'feed-post-link'})
    
    #titulo da noticia
    documento.add_paragraph("Titulo:")
    run = documento.add_paragraph()
    tituloword = run.add_run(titulo)
    tituloword.font.color.rgb = RGBColor(255, 0, 0)

    #subtitulo da noticia
    
    subtitulo = noticia.find('div', attrs={'class': 'feed-post-body-resumo'})
    if None != subtitulo:
      documento.add_paragraph("Subtitulo:")
      run2 = documento.add_paragraph()
      subtituloword = run2.add_run(subtitulo)
      subtituloword.font.color.rgb = RGBColor(0,255,0)
    
    #link da noticia
    documento.add_paragraph("Link:")
    run3 = documento.add_paragraph()
    linkword = run3.add_run(titulo['href'])
    linkword.font.color.rgb = RGBColor(0, 191, 255)

    documento.save('NoticiasDeHoje.doc')


def integracaoComOutolook():

    # integração com o outlook
    outlook = win32.Dispatch('outlook.application')

    # criando e-mail
    email = outlook.CreateItem(0)

    nome = 'Jocyanno'

    # configurações das informações do e-mail
    email.To = "jocyannovittor@hotmail.com"
    email.Subject = "Noticias Atualizadas"
    email.HTMLBody = f"""
    <p>Olá {nome}, aqui é o código Python Automático</p>
    <p>O arquivo com noticias atualizadas segue em anexo</p>
    <p>Boa Leitura!</p>

    """

    anexo = "C://Users/Jocya/Desktop/Nova_pasta/NoticiasDeHoje.doc"
    email.Attachments.Add(anexo)

    email.Send()
    print("Email Enviado")

#schedule.cada.tempo.fazer
schedule.every().day.at("10:57").do(webscrabing)
schedule.every().day.at("10:57").do(fazerColocarNoticiasNoWord)
schedule.every().day.at("10:57").do(integracaoComOutolook)

while 1:
      schedule.run_pending()
      time.sleep(1)
